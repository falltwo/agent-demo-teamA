from __future__ import annotations

from dataclasses import dataclass
from difflib import HtmlDiff
from io import BytesIO
import logging
from pathlib import Path
from typing import Any

from pypdf import PdfReader

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocument:
    name: str
    source: str
    text: str
    parser: str
    used_ocr: bool
    page_count: int | None = None
    warnings: tuple[str, ...] = ()


@dataclass(frozen=True)
class DiffSummary:
    left_name: str
    right_name: str
    html: str
    changed_lines: int
    added_lines: int
    removed_lines: int


def _coerce_name(source: str | Path | Any) -> str:
    if isinstance(source, Path):
        return source.name
    return str(getattr(source, "name", source))


def _extract_text_from_docx_bytes(raw: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(raw))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_text_from_pdf_bytes(raw: bytes) -> tuple[str, PdfReader]:
    reader = PdfReader(BytesIO(raw))
    texts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            texts.append(page_text.strip())
    return "\n\n".join(texts), reader


def _collect_pdf_page_images(reader: PdfReader) -> list[tuple[str, bytes]]:
    images: list[tuple[str, bytes]] = []
    for page_idx, page in enumerate(reader.pages):
        page_images = getattr(page, "images", None) or []
        for image_idx, image in enumerate(page_images):
            data = getattr(image, "data", None)
            if not data:
                continue
            ext = getattr(image, "image_format", None) or getattr(image, "extension", None) or "png"
            images.append((f"page-{page_idx + 1}-image-{image_idx + 1}.{ext}", data))
    return images


def _ocr_images_with_gemini(
    *,
    chat_client: Any,
    model: str,
    images: list[tuple[str, bytes]],
) -> str:
    try:
        from google.genai import types
    except Exception:
        return ""

    extracted: list[str] = []
    prompt = (
        "請將這張合約掃描圖完整 OCR 成純文字，保留段落、條號、金額、日期與甲乙方資訊。"
        "不要摘要，不要補寫看不到的內容。"
    )
    for filename, data in images:
        suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
        mime_type = {
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "png": "image/png",
            "bmp": "image/bmp",
            "gif": "image/gif",
            "webp": "image/webp",
            "tif": "image/tiff",
            "tiff": "image/tiff",
        }.get(suffix, "image/png")
        try:
            response = chat_client.models.generate_content(
                model=model,
                contents=[
                    prompt,
                    types.Part.from_bytes(data=data, mime_type=mime_type),
                ],
            )
            text = (getattr(response, "text", None) or "").strip()
            if text:
                extracted.append(text)
        except Exception as exc:
            logger.warning("OCR failed for %s: %s", filename, exc)
    return "\n\n".join(extracted)


def parse_uploaded_document(
    *,
    uploaded_file: Any,
    source: str,
    chat_client: Any | None = None,
    ocr_model: str | None = None,
    enable_ocr: bool = True,
) -> ParsedDocument | None:
    name = _coerce_name(uploaded_file)
    lower_name = name.lower()
    raw = uploaded_file.getvalue()
    warnings: list[str] = []

    if lower_name.endswith(".pdf"):
        try:
            text, reader = _extract_text_from_pdf_bytes(raw)
        except Exception as exc:
            logger.warning("PDF parse failed for %s: %s", name, exc, exc_info=True)
            return None

        used_ocr = False
        parser = "pdf_text"
        if enable_ocr and len(text.strip()) < 80 and chat_client is not None and ocr_model:
            images = _collect_pdf_page_images(reader)
            if images:
                ocr_text = _ocr_images_with_gemini(chat_client=chat_client, model=ocr_model, images=images)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    used_ocr = True
                    parser = "pdf_ocr"
            else:
                warnings.append("PDF 文字層不足，且未找到可供 OCR 的嵌入影像。")
        elif len(text.strip()) < 80:
            warnings.append("PDF 文字層不足，未啟用 OCR 或缺少 OCR 模型。")

        if not text.strip():
            return None
        return ParsedDocument(
            name=name,
            source=source,
            text=text,
            parser=parser,
            used_ocr=used_ocr,
            page_count=len(reader.pages),
            warnings=tuple(warnings),
        )

    if lower_name.endswith(".docx"):
        try:
            text = _extract_text_from_docx_bytes(raw)
        except Exception as exc:
            logger.warning("DOCX parse failed for %s: %s", name, exc, exc_info=True)
            return None
        if not text.strip():
            return None
        return ParsedDocument(
            name=name,
            source=source,
            text=text,
            parser="docx_text",
            used_ocr=False,
        )

    if lower_name.endswith(".txt") or lower_name.endswith(".md"):
        text = raw.decode("utf-8", errors="ignore")
        if not text.strip():
            return None
        return ParsedDocument(
            name=name,
            source=source,
            text=text,
            parser="plain_text",
            used_ocr=False,
        )

    return None


def parse_path_document(
    *,
    path: Path,
    chat_client: Any | None = None,
    ocr_model: str | None = None,
    enable_ocr: bool = True,
) -> ParsedDocument | None:
    class _PathFile:
        def __init__(self, file_path: Path):
            self.name = file_path.name
            self._raw = file_path.read_bytes()

        def getvalue(self) -> bytes:
            return self._raw

    return parse_uploaded_document(
        uploaded_file=_PathFile(path),
        source=str(path).replace("\\", "/"),
        chat_client=chat_client,
        ocr_model=ocr_model,
        enable_ocr=enable_ocr,
    )


def build_contract_diff(*, left_name: str, left_text: str, right_name: str, right_text: str) -> DiffSummary:
    left_lines = left_text.splitlines()
    right_lines = right_text.splitlines()
    changed_lines = 0
    added_lines = 0
    removed_lines = 0
    max_len = max(len(left_lines), len(right_lines))
    for idx in range(max_len):
        left_line = left_lines[idx] if idx < len(left_lines) else ""
        right_line = right_lines[idx] if idx < len(right_lines) else ""
        if left_line == right_line:
            continue
        changed_lines += 1
        if left_line and not right_line:
            removed_lines += 1
        elif right_line and not left_line:
            added_lines += 1
        else:
            added_lines += 1
            removed_lines += 1

    html = HtmlDiff(wrapcolumn=80).make_table(
        left_lines,
        right_lines,
        fromdesc=left_name,
        todesc=right_name,
        context=True,
        numlines=2,
    )
    return DiffSummary(
        left_name=left_name,
        right_name=right_name,
        html=html,
        changed_lines=changed_lines,
        added_lines=added_lines,
        removed_lines=removed_lines,
    )
