from io import BytesIO

from docx import Document

from document_processing import build_contract_diff, parse_uploaded_document


class _UploadedFile:
    def __init__(self, name: str, raw: bytes):
        self.name = name
        self._raw = raw

    def getvalue(self) -> bytes:
        return self._raw


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Master Service Agreement")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Party"
    table.rows[0].cells[1].text = "ACME"
    table.rows[1].cells[0].text = "Term"
    table.rows[1].cells[1].text = "12 months"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_uploaded_document_plain_text():
    uploaded = _UploadedFile("contract.txt", "Payment due in 30 days.".encode("utf-8"))
    parsed = parse_uploaded_document(uploaded_file=uploaded, source="uploaded/contract.txt", enable_ocr=False)

    assert parsed is not None
    assert parsed.parser == "plain_text"
    assert parsed.text == "Payment due in 30 days."
    assert not parsed.used_ocr


def test_parse_uploaded_document_docx_includes_table_text():
    uploaded = _UploadedFile("contract.docx", _make_docx_bytes())
    parsed = parse_uploaded_document(uploaded_file=uploaded, source="uploaded/contract.docx", enable_ocr=False)

    assert parsed is not None
    assert parsed.parser == "docx_text"
    assert "Master Service Agreement" in parsed.text
    assert "Party | ACME" in parsed.text
    assert "Term | 12 months" in parsed.text


def test_build_contract_diff_reports_line_changes():
    summary = build_contract_diff(
        left_name="v1.txt",
        left_text="Line A\nLine B\nLine C",
        right_name="v2.txt",
        right_text="Line A\nLine B updated\nLine C\nLine D",
    )

    assert summary.changed_lines >= 2
    assert summary.added_lines >= 1
    assert summary.removed_lines >= 1
    assert "v1.txt" in summary.html
    assert "v2.txt" in summary.html
